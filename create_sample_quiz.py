"""
Generate Sample Quiz Word Document
This script creates a sample quiz document with 10 questions in the required format.
Run this to generate sample_quiz.docx for testing the Endiba Quiz application.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_quiz_document():
    """Create a sample quiz Word document with 10 questions"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Endiba Quiz - Sample Questions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This is a sample quiz document for testing the Endiba Quiz application.')
    doc.add_paragraph('Each question has 4 options and the correct answer is indicated at the end.')
    doc.add_paragraph('')
    
    # Questions
    questions = [
        {
            'number': 1,
            'body': 'What is the capital city of France?',
            'options': {
                'A': 'London',
                'B': 'Paris',
                'C': 'Berlin',
                'D': 'Madrid'
            },
            'answer': 'B'
        },
        {
            'number': 2,
            'body': 'Which programming language is known for its use in web development and has a famous "Python" named after it?',
            'options': {
                'A': 'Java',
                'B': 'C++',
                'C': 'Python',
                'D': 'Ruby'
            },
            'answer': 'C'
        },
        {
            'number': 3,
            'body': 'What is the largest planet in our solar system?',
            'options': {
                'A': 'Saturn',
                'B': 'Jupiter',
                'C': 'Neptune',
                'D': 'Uranus'
            },
            'answer': 'B'
        },
        {
            'number': 4,
            'body': 'In what year did World War II end?',
            'options': {
                'A': '1942',
                'B': '1945',
                'C': '1948',
                'D': '1950'
            },
            'answer': 'B'
        },
        {
            'number': 5,
            'body': 'What is the chemical symbol for gold?',
            'options': {
                'A': 'Go',
                'B': 'Gd',
                'C': 'Au',
                'D': 'Ag'
            },
            'answer': 'C'
        },
        {
            'number': 6,
            'body': 'Which scientist developed the theory of relativity?',
            'options': {
                'A': 'Isaac Newton',
                'B': 'Albert Einstein',
                'C': 'Niels Bohr',
                'D': 'Stephen Hawking'
            },
            'answer': 'B'
        },
        {
            'number': 7,
            'body': 'What is the smallest prime number?',
            'options': {
                'A': '0',
                'B': '1',
                'C': '2',
                'D': '3'
            },
            'answer': 'C'
        },
        {
            'number': 8,
            'body': 'Which continent is known as the "Land of the Rising Sun"?',
            'options': {
                'A': 'China',
                'B': 'India',
                'C': 'Japan',
                'D': 'Thailand'
            },
            'answer': 'C'
        },
        {
            'number': 9,
            'body': 'What is the largest ocean on Earth?',
            'options': {
                'A': 'Atlantic Ocean',
                'B': 'Indian Ocean',
                'C': 'Arctic Ocean',
                'D': 'Pacific Ocean'
            },
            'answer': 'D'
        },
        {
            'number': 10,
            'body': 'Which organ in the human body pumps blood throughout the circulatory system?',
            'options': {
                'A': 'Brain',
                'B': 'Liver',
                'C': 'Heart',
                'D': 'Lungs'
            },
            'answer': 'C'
        }
    ]
    
    for q in questions:
        # Add question
        doc.add_paragraph(f'Question {q["number"]}: {q["body"]}')
        
        # Add options
        for label, text in q['options'].items():
            doc.add_paragraph(f'{label}. {text}')
        
        # Add answer
        doc.add_paragraph(f'Question {q["number"]} Answer: {q["answer"]}')
        doc.add_paragraph('')  # Empty line between questions
    
    # Save document
    doc.save('sample_quiz.docx')
    print('Sample quiz document created: sample_quiz.docx')
    print('This file contains 10 sample questions for testing Endiba Quiz.')


if __name__ == '__main__':
    create_sample_quiz_document()
